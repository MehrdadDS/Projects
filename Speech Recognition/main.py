import speech_recognition as sr
from docx import Document
from pydub import AudioSegment

# Initialize the recognizer
recognizer = sr.Recognizer()

# Path to the audio file
audio_file_path = r"C:\My Folder\Github\Speech Recognition\20240623_115021_Weston Station 2.wav"
converted_audio_file_path = r"C:\My Folder\Github\Speech Recognition\converted_audio.wav"

# Convert audio to WAV format if necessary
try:
    audio = AudioSegment.from_file(audio_file_path)
    audio.export(converted_audio_file_path, format="wav")
    print("Audio file converted successfully.")
except Exception as e:
    print(f"Error converting audio file: {e}")
    exit()

# Load the audio file
with sr.AudioFile(converted_audio_file_path) as source:
    # Listen to the audio file
    audio_data = recognizer.record(source)

# Convert the audio to text
try:
    text = recognizer.recognize_google(audio_data)
    print("Transcription: " + text)

    # Save the transcription to a Word file
    document = Document()
    document.add_heading('Transcription', level=1)
    document.add_paragraph(text)

    # Path to save the Word file
    word_file_path = r"C:\My Folder\Github\Speech Recognition\Transcription.docx"
    document.save(word_file_path)
    print(f"Transcription saved to {word_file_path}")

except sr.UnknownValueError:
    print("Google Speech Recognition could not understand the audio")
except sr.RequestError as e:
    print(f"Could not request results from Google Speech Recognition service; {e}")
